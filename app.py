import streamlit as st
import pandas as pd
import numpy as np
import io
import json
import xlsxwriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
import re

# ==========================================
# PART 1: 配置区域
# ==========================================

COMMON_METRICS = {
    "spend": ["花费金额(USD)", "花费金额", "Amount Spent", "Cost", "花费"],
    "roas": ["广告花费回报 (ROAS) - 购物", "ROAS", "Purchase ROAS", "Return on Ad Spend"],
    "purchases": ["购买次数", "成效", "Purchases", "Results", "Website Purchases"],
    "cpa": ["单次购买费用", "Cost per Purchase", "Cost per Result", "CPA"],
    "ctr": ["链接点击率", "CTR", "Link CTR"],
    "cpm": ["千次展示费用", "CPM"],
    "clicks": ["点击", "链接点击", "Clicks", "Link Clicks"],
    "impressions": ["曝光", "展示次数", "Impressions"],
    "purchase_value": ["购买价值", "购物价值", "Purchase Value", "Conversion Value", "Total Value"],
    "aov": ["单次购买价值", "单次购物价值"]
}

SHEET_MAPPINGS = {
    "整体数据": {
        **COMMON_METRICS,
        "date_range": ["时间范围", "Date Range", "Time"],
        "clicks_all": ["点击", "点击(全部)", "Clicks (All)"],
        "landing_page_views": ["落地页浏览量", "落地页", "Landing Page Views", "Landing", "落地页浏览"],
        # ✅ 这里的定义只做参考，实际逻辑在 find_column_smart 中加强
        "add_to_cart": ["加入购物车", "加购", "Add to Cart", "Website Adds to Cart", "Adds to Cart", "Cart"], 
        "initiate_checkout": ["结账发起次数", "结账", "Initiate Checkout", "Checkouts"],
        "rate_click_to_lp": ["点击-落地页浏览转化率"],
        "rate_lp_to_atc": ["落地页浏览-加购转化率"],
        "rate_atc_to_ic": ["加购-结账转化率"],
        "rate_ic_to_pur": ["结账-购买转化率"]
    },
    "分时段数据": {
        **COMMON_METRICS,
        "date_range": ["时间范围", "Day", "Date", "Reporting Starts"],
        "landing_page_views": ["落地页浏览量", "Landing Page Views"],
        "add_to_cart": ["加入购物车", "加购", "Add to Cart"],
        "initiate_checkout": ["结账发起次数", "Initiate Checkout"],
    },
    "异常指标": {"anomaly_metric_name": ["异常指标"], "mom_change": ["环比"]},
    "广告架构": {**COMMON_METRICS, "dimension_item": ["广告类型"]},
    "受众组": {
        **COMMON_METRICS,
        "dimension_item": ["广告组", "广告组Id", "Ad Set Name"],
        "custom_audience_settings": ["设置的自定义受众", "Custom Audiences"],
        "converting_keywords": ["产生成效的关键词", "Interests"],
        "converting_countries": ["产生成效的国家", "国家", "Country"],
        "converting_genders": ["产生成效的性别", "性别", "Gender"],
        "converting_ages": ["产生成效的年龄", "年龄", "Age"]
    },
    "受众类型": {**COMMON_METRICS, "dimension_item": ["受众类型"]},
    "国家": {**COMMON_METRICS, "dimension_item": ["国家/地区", "国家"]},
    "年龄": {**COMMON_METRICS, "dimension_item": ["年龄"]},
    "性别": {**COMMON_METRICS, "dimension_item": ["性别"]},
    "平台&版位": {**COMMON_METRICS, "dimension_item": ["平台&版位"]},
    "素材": {
        **COMMON_METRICS,
        "content_item": ["素材", "Ad Name", "Creative Name"],
        "cvr_lp_to_pur": ["落地页浏览-购买转化率"]
    },
    "落地页": {
        **COMMON_METRICS,
        "content_item": ["落地页url", "落地页", "Website URL"],
        "ctr_all": ["曝光-点击转化率"],
        "rate_lp_to_atc": ["落地页浏览-加购转化率"]
    }
}

GROUP_CONFIG = {
    "Master_Overview": ["整体数据", "分时段数据", "异常指标"],
    "Master_Breakdown": ["广告架构", "受众组", "受众类型", "国家", "年龄", "性别", "平台&版位"],
    "Master_Creative": ["素材", "落地页"]
}

REPORT_MAPPING = {
    "spend": "花费 ($)", "roas": "ROAS", "purchases": "购买次数", "purchase_value": "购买总价值",
    "cpa": "CPA ($)", "ctr": "CTR (%)", "cpm": "CPM ($)", "aov": "客单价",
    "impressions": "展现量", "clicks_all": "点击量 (All)", "clicks": "点击量 (All)", "ctr_all": "点击率 (All)",
    "landing_page_views": "落地页访问量", "add_to_cart": "加购次数", "initiate_checkout": "结账发起数 (IC)",
    "rate_click_to_lp": "点击 → 落地页访问转化率", "rate_lp_to_atc": "落地页 → 加购转化率",
    "rate_atc_to_ic": "加购 → 购买转化率", "rate_ic_to_pur": "购买转化率",
    "cvr_purchase": "点击 → 购买转化率", "cvr_lp_to_pur": "CVR (全站转化率)",
    "date_range": "日期/时段", "campaign_type": "投放模式", "adset_name": "广告组ID",
    "custom_audience_settings": "自定义受众源", "converting_keywords": "高潜兴趣词", 
    "country": "国家", "creative_name": "素材名称", "placement": "版位",
    "landing_page_url": "页面 URL", "mom_change": "环比波动", "anomaly_metric_name": "异常项",
    "converting_countries": "产生成效的国家", "converting_genders": "产生成效的性别", "converting_ages": "产生成效的年龄"
}

FIELD_ALIASES = {
    "spend": ["spend", "cost", "花费"],
    "purchases": ["purchases", "results", "成效", "购买"],
    "clicks": ["clicks", "点击"],
    "impressions": ["impressions", "展示"],
    "add_to_cart": ["add_to_cart", "cart", "加购"],
    "initiate_checkout": ["initiate_checkout", "checkout", "结账"],
    "landing_page_views": ["landing_page_views", "落地页", "landing"]
}

# ==========================================
# PART 2: 核心工具函数
# ==========================================

def clean_numeric_strict(val): 
    if pd.isna(val): return 0.0
    if isinstance(val, (int, float)): return float(val)
    # 处理特殊字符
    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')
    if val_str == '-' or val_str == '—': return 0.0 # 处理 Excel 里的横杠
    if '%' in val_str: 
        val_str = val_str.replace('%', '')
        try: return float(val_str) / 100.0
        except: return 0.0
    try: return float(val_str)
    except: return 0.0

def clean_numeric(val):
    # 与 strict 类似，但用于 DataFrame apply，宽容度高一点
    if pd.isna(val): return 0.0
    if isinstance(val, (int, float)): return float(val)
    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')
    if val_str == '-' or val_str == '—': return 0.0
    if '%' in val_str: 
        val_str = val_str.replace('%', '')
        try: return float(val_str) / 100.0 
        except: return 0.0
    try: return float(val_str)
    except: return val 

def safe_div(n, d, m=1.0):
    n_val, d_val = clean_numeric_strict(n), clean_numeric_strict(d)
    return (n_val / d_val * m) if d_val > 0 else 0.0

def find_column_smart(df, target_key, keywords):
    """
    智能列名匹配：
    1. 优先完全匹配
    2. 其次模糊匹配（包含关键词）
    3. ✅ 核心修正：排除 'Cost', 'Value', 'Rate' 等干扰词，防止把 'Cost per Add to Cart' 识别为 'Add to Cart'
    """
    # 排除词列表：如果目标是计数类指标（加购、结账、购买），不能包含这些词
    exclusion_list = []
    if target_key in ['add_to_cart', 'initiate_checkout', 'purchases', 'clicks', 'impressions']:
        exclusion_list = ['cost', 'cpa', 'value', 'rate', '费用', '价值', '率', '单次']
    
    # 1. 优先级最高：全字匹配 (Case Insensitive)
    for col in df.columns:
        for kw in keywords:
            if kw.lower() == col.lower():
                return col

    # 2. 优先级第二：包含匹配 (但要检查排除词)
    for col in df.columns:
        col_lower = col.lower()
        # 必须包含关键词
        is_match = False
        for kw in keywords:
            if kw.lower() in col_lower:
                is_match = True
                break
        
        if is_match:
            # 检查是否包含排除词
            has_exclusion = False
            for exc in exclusion_list:
                if exc in col_lower:
                    has_exclusion = True
                    break
            
            if not has_exclusion:
                return col
    
    return None

def calc_metrics_dict(df_chunk):
    res = {}
    if df_chunk.empty: return res
    sums = {}
    targets = ['spend', 'clicks', 'impressions', 'purchases', 'purchase_value',
               'landing_page_views', 'add_to_cart', 'initiate_checkout']
    
    for t in targets:
        col = find_column_smart(df_chunk, t, FIELD_ALIASES.get(t, [t]))
        if col:
             sums[t] = df_chunk[col].apply(clean_numeric_strict).sum()
        else:
             sums[t] = 0.0

    res.update(sums)
    res['roas'] = safe_div(sums.get('purchase_value'), sums.get('spend'))
    res['cpm'] = safe_div(sums.get('spend'), sums.get('impressions'), 1000)
    res['cpc'] = safe_div(sums.get('spend'), sums.get('clicks'))
    res['ctr'] = safe_div(sums.get('clicks'), sums.get('impressions'))
    res['cpa'] = safe_div(sums.get('spend'), sums.get('purchases'))
    res['cvr_purchase'] = safe_div(sums.get('purchases'), sums.get('clicks'))
    res['rate_click_to_lp'] = safe_div(sums.get('landing_page_views'), sums.get('clicks'))
    res['rate_lp_to_atc']   = safe_div(sums.get('add_to_cart'), sums.get('landing_page_views'))
    res['rate_atc_to_ic']   = safe_div(sums.get('initiate_checkout'), sums.get('add_to_cart'))
    res['rate_ic_to_pur']   = safe_div(sums.get('purchases'), sums.get('initiate_checkout'))
    res['aov'] = safe_div(sums.get('purchase_value'), sums.get('purchases'))
    return res 

def format_cell(key, val, is_mom=False):
    if isinstance(val, str): return val
    if is_mom: return val if key == 'date_range' else f"{val:+.2%}"
    k = str(key).lower()
    if 'roas' in k: return f"{val:.2f}"
    if any(x in k for x in ['rate', 'ctr', 'cvr']): return f"{val:.2%}" 
    if any(x in k for x in ['spend', 'cpm', 'cpc', 'value', 'aov', 'cpa']): return f"{val:,.2f}"
    if any(x in k for x in ['purchases', 'cart', 'click', 'impressions', 'checkout']): return f"{val:,.0f}"
    return f"{val}"

def add_df_to_word(doc, df, title, level=1):
    if df.empty: return
    doc.add_heading(title, level=level)
    t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    t.style = 'Table Grid'
    for j, col in enumerate(df.columns):
        cell = t.cell(0, j)
        cell.text = str(col)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = df.iat[i, j]
            cell = t.cell(i+1, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    doc.add_paragraph("\n")

# ==========================================
# PART 3: 主逻辑类 (ETL + 调试诊断)
# ==========================================

class AdReportProcessor:
    def __init__(self, raw_file, bench_file=None):
        self.raw_file = raw_file
        self.bench_file = bench_file
        self.processed_dfs = {}
        self.merged_dfs = {}
        self.final_json = {}
        self.doc = Document()
        self.debug_log = [] # 专门用于前端显示的诊断日志

    def find_sheet_fuzzy(self, target, actual_sheets):
        for actual in actual_sheets:
            if target.strip().lower() == actual.strip().lower(): return actual
        for actual in actual_sheets:
            if target in actual: return actual
        return None

    def process_etl(self):
        xls = pd.ExcelFile(self.raw_file)
        
        for config_sheet_name, mapping in SHEET_MAPPINGS.items():
            actual_sheet_name = self.find_sheet_fuzzy(config_sheet_name, xls.sheet_names)
            
            if actual_sheet_name:
                df = pd.read_excel(xls, sheet_name=actual_sheet_name)
                df.columns = [str(c).strip() for c in df.columns]
                
                final_cols = {}
                # ✅ 1. 智能匹配列
                for std_col, raw_col_options in mapping.items():
                    # 结合配置的别名 + 智能排除逻辑
                    search_keywords = raw_col_options
                    matched_col = find_column_smart(df, std_col, search_keywords)
                    
                    if matched_col:
                        final_cols[std_col] = matched_col
                        # 📝 记录关键指标的匹配情况到日志
                        if config_sheet_name == "整体数据" and std_col in ["add_to_cart", "initiate_checkout"]:
                            sample_vals = df[matched_col].head(3).tolist()
                            self.debug_log.append(f"✅ [整体数据] 成功匹配 '{std_col}' -> 原列名 '{matched_col}' | 前3行数据: {sample_vals}")
                
                # ✅ 2. 兜底匹配：如果没找到加购，尝试暴力搜索包含 "cart" 且不含 "cost" 的列
                if config_sheet_name == "整体数据" and "add_to_cart" not in final_cols:
                    for c in df.columns:
                        if "cart" in c.lower() and "cost" not in c.lower() and "value" not in c.lower():
                            final_cols["add_to_cart"] = c
                            self.debug_log.append(f"⚠️ [整体数据] 暴力匹配 '{c}' 为 add_to_cart")
                            break

                # 3. 创建 Clean DF
                if final_cols:
                    df_clean = df[list(final_cols.values())].rename(columns={v: k for k, v in final_cols.items()})
                else:
                    df_clean = pd.DataFrame()
                
                # 4. 补全缺失列为 0
                for expected_col in mapping.keys():
                    if expected_col not in df_clean.columns:
                        df_clean[expected_col] = 0.0
                        if config_sheet_name == "整体数据" and expected_col == "add_to_cart":
                            self.debug_log.append(f"❌ [整体数据] 未找到 '加购' 相关列，已填充为0。请检查原表列名。")

                # 5. 数
